from docx import Document
from io import BytesIO
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_resume(name, email, phone, linkedin, github, education, experience, skills,
                    career_objective, photo, summary, certifications, projects, achievements, languages):
    doc = Document()

    if photo:
        doc.add_picture(photo, width=Inches(1.5))
        doc.add_paragraph()  # Add space after photo

    doc.add_heading(name, 0)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = contact.add_run(f"📧 {email} | 📞 {phone}\n🔗 LinkedIn: {linkedin} | 💻 GitHub: {github}")
    run.font.size = Pt(10)

    def styled_heading(text):
        heading = doc.add_heading(level=1)
        run = heading.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if summary:
        styled_heading("Profile Summary")
        doc.add_paragraph(summary)

    if career_objective:
        styled_heading("Career Objective")
        doc.add_paragraph(career_objective)

    if education:
        styled_heading("Education")
        doc.add_paragraph(education)

    if experience:
        styled_heading("Work Experience")
        doc.add_paragraph(experience)

    if skills:
        styled_heading("Skills")
        doc.add_paragraph("\n".join(f"• {skill.strip()}" for skill in skills.split(",") if skill.strip()))

    if certifications:
        styled_heading("Certifications")
        doc.add_paragraph(certifications)

    if projects:
        styled_heading("Projects")
        doc.add_paragraph(projects)

    if achievements:
        styled_heading("Achievements / Awards")
        doc.add_paragraph(achievements)

    if languages:
        styled_heading("Languages Known")
        doc.add_paragraph("\n".join(f"• {lang.strip()}" for lang in languages.split(",") if lang.strip()))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
